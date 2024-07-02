import os
from docx import Document
from docx.shared import RGBColor, Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
import matplotlib.pyplot as plt
from matplotlib.font_manager import FontProperties

from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase import pdfmetrics


def collect_file_info(directory, file_type):  # 收集文件信息
    """
    :param directory: 文件根路径
    :param file_type: 所需文件类型
    :return: 所需文件的路径
    """
    file_list = []
    for root, dirs, files in os.walk(directory):
        for file in files:
            if file.endswith(file_type):
                file_path = os.path.join(root, file)
                file_list.append(file_path)  # 添加当前目录下的所有文件路径
        for dir in dirs:
            _ = collect_file_info(dir, file_type)
            file_list += _  # 添加子目录下的所有文件路径
    return file_list


def identify_and_extract_file_contents(file_list, save_name):  # 识别和提取文件内容
    """
    :param file_list: 需要识别和提取的所有文件列表
    :param save_name: 需要保存的Word文档名称
    :return: 无返回
    """
    document = Document()  # 创建Word文档
    for file_path in file_list:
        with open(file_path, 'r', encoding='utf-8') as f:
            news = f.readlines()  # 读取文件
            if '标题' not in news[0] or '日期' not in news[1] or '内容' not in news[2]:  # 判断是否是新闻文件
                continue  # 不是新闻文件直接跳过
            count = 0
            while count < len(news):
                if '标题' in news[count]:  # 找到标题
                    title = document.add_heading("", level=1)  # 设置1级标题
                    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = title.add_run(news[count].split('标题：')[1])
                    run.font.name = u'宋体'  # 设置标题字体为宋体
                    run.font.color.rgb = RGBColor(0, 0, 0)  # 设置标题颜色为黑色
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')

                    count += 1  # 找到日期并添加
                    date = document.add_heading('', level=3)
                    date.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = date.add_run(news[count].split('日期：')[1])
                    run.font.name = u'Times New Roman'  # 设置日期字体为Times New Roman
                    run.font.color.rgb = RGBColor(0, 0, 0)  # 设置日期颜色为黑色
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'Times New Roman')

                    count += 2  # 找到内容并添加
                    while count < len(news):
                        if news[count].strip() == '':  # 空行是不同新闻间区分
                            count += 1
                            break
                        content = document.add_paragraph(news[count].strip())  # 添加一段新闻内容
                        content.paragraph_format.line_spacing = 1.5  # 1.5倍行间距
                        content.paragraph_format.first_line_indent = Pt(22)  # 首行缩进22磅（两个字符）
                        content.paragraph_format.space_before = Pt(0)  # 段前0磅
                        content.paragraph_format.space_after = Pt(0)  # 段后0磅
                        for run in content.runs:
                            run.font.name = u'Times New Roman'  # 设置西文为Times New Roman
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')  # 设置中文为宋体
                            run.font.size = Pt(11)  # 设置字体大小为11磅
                        count += 1

    document.save(save_name)  # 保存Word文档，并命名


def set_run(run, font_size, color, highlight_color):  # 定义设置文本样式的函数set_run，按需添加
    run.font.size = font_size
    run.font.color.rgb = color
    run.font.highlight_color = highlight_color
    run.font.name = u'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')


def parse_file(file_path, key_words):
    """
    :param file_path: 需要解析的Word文档
    :key_words: 关键词文件
    :return: 关键词字典keywords
    """
    with open(key_words, 'r', encoding='utf-8') as f:  # 提取关键词
        keywords = {}
        for line in f.readlines():
            keywords[line.strip()] = 0
    docx = Document(file_path)
    for keyword in keywords:
        for paragraph in docx.paragraphs:
            if keyword in paragraph.text:  # 检查段落中有无关键词
                for r in paragraph.runs:  # 遍历段落的所有格式
                    if r.text != '':  # 过滤空的run
                        # 记录之前的状态
                        font_size = r.font.size
                        color = r.font.color.rgb
                        highlight_color = r.font.highlight_color

                        rest = r.text.split(keyword)  # 以关键词做切分
                        keywords[keyword] += len(rest) - 1
                        r.text = ''  # 清空原有run文字
                        for text in rest[:-1]:
                            run = paragraph.add_run(text)
                            set_run(run, font_size, color, highlight_color)
                            run = paragraph.add_run(keyword)
                            set_run(run, font_size, color, WD_COLOR_INDEX.YELLOW)  # 修改成高亮
                        run = paragraph.add_run(rest[-1])
                        set_run(run, font_size, color, highlight_color)
    docx.save(file_path)
    return keywords


def convert_docx_to_pdf(docx_path, pdf_path):  # 将 Word 文档保存为 PDF
    docx = Document(docx_path)  # 打开Word文档
    pdf = SimpleDocTemplate(pdf_path, pagesize=letter)  # 新建一个PDF对象
    styles = getSampleStyleSheet()

    pdfmetrics.registerFont(TTFont('SimSun', 'SimSun.ttf'))  # 添加中文字体
    styles['Normal'].fontName = 'SimSun'

    content = []
    for p in docx.paragraphs:  # 将Word文件内容转为PDF
        content.append(Paragraph(p.text, styles['Normal']))

    pdf.build(content)


def visualize_keyword_occurrences(keyword_occurrences):  # 将关键词出现次数的统计结果可视化为柱状图
    keywords = list(keyword_occurrences.keys())
    counts = list(keyword_occurrences.values())

    with open('result.txt', 'w', encoding='utf-8') as f:
        for key, value in keyword_occurrences.items():
            f.write(f'{key}：{value}\n')

    # 指定中文字体
    font = FontProperties(fname=r'C:\Windows\Fonts\SimHei.ttf', size=12)

    plt.figure(figsize=(8, 6))
    plt.bar(keywords, counts, color='skyblue')
    plt.xlabel('关键词', fontproperties=font)
    plt.ylabel('出现次数', fontproperties=font)
    plt.title('文档中关键词出现次数', fontproperties=font)
    plt.xticks(rotation=45, ha='right', fontproperties=font)  # 旋转45度，右对齐
    plt.tight_layout()

    plt.show()


def main():
    # 文件信息收集
    # destination_path = input('请输入目标文件名称：')
    # file_type = input("请输入所需文件类型（目前只支持txt）：")
    destination_path = 'dir1'
    file_type = 'txt'
    file_list = collect_file_info(destination_path, file_type)

    # 文件识别与分析
    # key_words = input('请输入关键词字典：')
    # document = input('请输入要保存的Word文档名称：')
    key_words = 'key_words.txt'
    document = 'news.docx'
    identify_and_extract_file_contents(file_list, document)
    keyword_occurrences = parse_file(document, key_words)

    # 将Word转换为PDF
    # PDF = input('请输入要转换成PDF的名称：')
    PDF = 'news.pdf'
    convert_docx_to_pdf(document, PDF)

    # 可视化关键词出现次数
    visualize_keyword_occurrences(keyword_occurrences)


if __name__ == '__main__':
    main()
